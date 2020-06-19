import json
import os
import io
import time
import re
import json
from copy import deepcopy
from docx import Document
from docx.shared import Inches
from werkzeug.datastructures import FileStorage

from flask import Flask
from app import app
from config import Config
from app.utils.utils import (
    does_file_exists,
    getNow,
    getUUID,
)
from app.constants import (
    GeneralConstants,
)
from app.handlers import format_exception
from app.exceptions import (
    InvalidDocumentFomat,
    DocumentSavingError,
)

# File storage objects


class FileStorageObject(FileStorage):

    def __init__(self, path, file_name, content_type, stream=None):
        self.path = os.path.join(path)
        self.stream = open(path, "rb") if stream is None else stream
        self.file_name = file_name
        super().__init__(stream=self.stream,
                         filename=self.file_name,
                         content_type=content_type,)


class BaseFile():

    def __init__(self, storage_object=None, full_path=None, method=None):
        self.full_path = full_path
        self.read_method = method
        self.storage_object = storage_object

    @property
    def storage_object(self):
        if (self.__storage_object is None):
            self.__storage_object = open(self.full_path, self.read_method)
        return self.__storage_object

    @storage_object.setter
    def storage_object(self, storage_obj):
        self.__storage_object = storage_obj

    def read(self):
        if 'r' in self.read_method:
            self.body = self.storage_object.read()

    def save(self):
        self.storage_object.save()
    
    def set_mode(self, read_method):
        self.read_method = read_method
        self.close()
        self.storage_object = open(self.full_path, self.read_method)

    def write(self, data):
        self.storage_object.write(data)

    def close(self):
        self.storage_object.close()

# File objects


class File(BaseFile):

    def __init__(self, storage_object=None, name=None, extension=None, read_method=None):
        self.name = name
        self.extension = extension
        self.read_method = read_method
        super().__init__(storage_object, self.full_path, self.read_method)

    @property
    def name(self):
        return self.__name

    @name.setter
    def name(self, name):
        self.__name = name

    @property
    def extension(self):
        return self.__extension

    @extension.setter
    def extension(self, extension):
        self.__extension = extension

    @property
    def full_name(self):
        self.__full_name = self.name + "." + self.extension
        return self.__full_name

    @property
    def path(self):
        self.__path = Config.RENDERED_FILES_FOLDER + \
            self.full_name
        return self.__path

    @path.setter
    def path(self, path):
        self.__path = path

    @property
    def full_path(self):
        self.__full_path = GeneralConstants.UPLOADS_PATH + self.path
        return self.__full_path

    @full_path.setter
    def full_path(self, full_path):
        self.__full_path = full_path

    def save(self):
        if not isinstance(self.storage_object, io.BufferedReader):
            self.storage_object.save(self.full_path)
        if does_file_exists(self.full_path):
            app.logger.info('File is saved.')
        else:
            raise DocumentSavingError()


class GeneratedFile(File):

    """
        File class with the additional output fields.

    """
    def __init__(self, name, storage_object=None, read_method='rb', output_name=GeneralConstants.TEMPLATE_PREFIX, extension=None):
        self.storage_object = storage_object
        self.name = name
        self.output_name = output_name
        self.extension = self.form_file_extension() if extension is None else extension
        self.read_method = read_method

    @property
    def output_name(self):
        return self.__output_name

    @output_name.setter
    def output_name(self, output_name):
        self.__output_name = output_name

    @property
    def output_full_name(self):
        self.__output_full_name = self.output_name + "." + self.extension
        return self.__output_full_name

    @output_full_name.setter
    def output_full_name(self, output_full_name):
        self.__output_full_name = output_full_name

    def form_file_extension(self):
        if "." in self.storage_object.filename:
            return self.storage_object.filename.split('.')[1]
        else:
            raise InvalidDocumentFomat()
    
    @classmethod
    def get_obj_by_name(cls, file_name, output_name=GeneralConstants.TEMPLATE_PREFIX, extension=GeneralConstants.TEMPLATE_FILE_EXTENSION):
        """
            The class method for creating GeneratedFile obj from only file_name.
        """
        return cls(name=file_name, output_name=output_name, extension=extension)

    @classmethod
    def make_copy(cls, file_object):
        """
            The class method for creating the GeneratedFile obj copy.
        """
        return cls(name=file_object.name, output_name=file_object.output_name, extension=file_object.extension)

class TemplateFile(GeneratedFile):
    """
        Docx template file class.
    """

    def __init__(self, name=None, storage_object=None, output_name=GeneralConstants.TEMPLATE_PREFIX, extension=None):
        super().__init__(name=name, storage_object=storage_object, output_name=output_name, extension=extension)
        self.save()

    def save(self):
        if self.extension == GeneralConstants.TEMPLATE_FILE_EXTENSION:
            return super().save()
        else:
            raise InvalidDocumentFomat()


class HTMLFile(GeneratedFile):

    def __init__(self, name, read_method, data=None, storage_object=None, output_name=GeneralConstants.TEMPLATE_PREFIX, extension=GeneralConstants.HTML_EXTENSION):
        super().__init__(name, storage_object=storage_object, read_method=read_method,  output_name=output_name, extension=extension)
        self.data = data

    @property
    def path(self):
        self.__path = Config.TEMP_FOLDER + \
            self.name + "." + self.extension
        return self.__path

    @property
    def full_path(self):
        self.__full_path = GeneralConstants.UPLOADS_PATH + Config.TEMPLATES_TEMP_FOLDER + \
            self.name + "." + self.extension
        return self.__full_path

    def write(self):
        self.storage_object.write(self.data)


class JSONFile(GeneratedFile):

    def __init__(self, name, read_method, data=None, storage_object=None, output_name=GeneralConstants.TEMPLATE_PREFIX, extension=GeneralConstants.JSON_EXTENSION):
        super().__init__(name=name, read_method=read_method, storage_object=storage_object,
                         output_name=output_name, extension=extension)
        self.data = data

    @property
    def data(self):
        return self.__data

    @data.setter
    def data(self, data):
        if data is not None:
            self.__data = json.loads(data)
        with self.storage_object as f:
            json.dump(self.__data, f, indent=3)


class PdfFile(GeneratedFile):

    def __init__(self, name=None, storage_object=None, output_name=GeneralConstants.TEMPLATE_PREFIX, extension=None):
        super().__init__(name=name, storage_object=storage_object, output_name=output_name, extension=extension)
        self.save()

class DocxFile:
    """
        Docx document class.
    """

    def __init__(self, name=None, folder=None):
        self.name = name
        self.extension = GeneralConstants.TEMPLATE_FILE_EXTENSION
        self.folder = folder
        self.object = Document()

    @property
    def name(self):
        return self.__name

    @name.setter
    def name(self, name):
        if name is not None:
            self.__name = name
        else:
            self.__name = str(getUUID())

    @property
    def extension(self):
        return self.__extension

    @extension.setter
    def extension(self, extension):
        self.__extension = extension

    @property
    def full_name(self):
        return self.name + "." + self.extension

    @property
    def path(self):
        return self.folder + self.full_name

    def save(self):
        self.object.save(self.path)

    def add_paragraph(self, data):
        p = self.object.add_paragraph(data)

    def add_heading(self, data, level):
        heading = self.object.add_heading(data, level=level)

    def add_table(row_number, col_number, table_data):
        """
            The method for table adding.
            Input:
                row_number- a row number including header row
                col_number - a column nomber including header row
                table_data - a matrix (list of lists)
        """
        table = self.object.add_table(rows=1, cols=3)
        for row in range(0, row_number):
            for column in range(0, col_number):
                table.rows[row].cells[column].text = str(
                    table_data[row][column])
            table.add_row().cells

    def add_page_break(self):
        self.object.add_page_break()

    def getAllText(self):
        return self.__class__.getAllText(document=self.object)

    def getTableContent(self, table_number):
        return self.__class__.getTableContent(table_number, document=self.object)

    def getAllTablesContent(self, table_number):
        return self.__class__.getAllTablesContent(document=self.object)

    @classmethod
    def getAllText(cls, filename=None, document=None):
        if(document is None):
            document = Document(filename)
        fullText = []
        for para in document.paragraphs:
            fullText.append(para.text)
        return '\n'.join(fullText)

    @classmethod
    def getTableContent(cls, table_number, filename=None, document=None):
        if(document is None):
            document = Document(filename)
        if table < table_number:
            table = document.tables[table_number]
        data = []
        keys = None
        for i, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)
            if i == 0:
                keys = tuple(text)
                continue
            row_data = dict(zip(keys, text))
            data.append(row_data)
        return data

    @classmethod
    def getAllTablesContent(cls, filename=None, document=None):
        if(document is None):
            document = Document(filename)
        data = []
        for table_number in range(0, len(document.tables)):
            table_data = cls.getTableContent(table_number, document=document)
            data.append(table_data)
        return data